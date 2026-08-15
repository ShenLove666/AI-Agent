from __future__ import annotations

import asyncio
import csv
import io
import os
import re
import threading
from collections.abc import Iterable, Iterator
from datetime import date, datetime, time
from pathlib import Path
from typing import Any

from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from app.framework.errors import AppError
from app.modules.knowledge.models import KnowledgeBase, KnowledgeChunk, KnowledgeDocument
from app.modules.vector.indexer import VectorIndexer


_POLICY_TERMS = ("政策", "规则", "退货", "退款", "售后")
_RECOMMENDATION_TERMS = ("指南", "推荐", "搭配")
_PRODUCT_TERMS = ("商品", "说明")

DEFAULT_MAX_PARSE_FILE_BYTES = 50 * 1024 * 1024
DEFAULT_MAX_PARSE_ROWS = 10_000
DEFAULT_MAX_PARSE_COLUMNS = 100
DEFAULT_MAX_PARSE_CELL_CHARS = 10_000
DEFAULT_MAX_PARSE_TEXT_CHARS = 2_000_000
DEFAULT_MAX_PARSE_SHEETS = 20
_CSV_FIELD_LIMIT_LOCK = threading.Lock()


def _env_limit(names: tuple[str, ...], default: int) -> int:
    """Read a positive parser limit, falling back safely on bad env values."""

    for name in names:
        raw = os.getenv(name)
        if raw is None:
            continue
        try:
            value = int(raw)
        except (TypeError, ValueError):
            continue
        if value > 0:
            return value
    return default


def infer_source_kind(*names: str) -> str:
    """按文件名/知识库名做初始 source_kind 推断（简单规则）。

    含 政策/规则/退货/退款/售后 → policy；
    含 指南/推荐/搭配 → recommendation_guide；
    含 商品/说明 → product_knowledge；
    其余 → general。
    """
    text = " ".join(name or "" for name in names)
    if any(term in text for term in _POLICY_TERMS):
        return "policy"
    if any(term in text for term in _RECOMMENDATION_TERMS):
        return "recommendation_guide"
    if any(term in text for term in _PRODUCT_TERMS):
        return "product_knowledge"
    return "general"


class DocumentParser:
    """Parse supported knowledge files into bounded plain text.

    Spreadsheet parsing deliberately uses stdlib ``csv`` and openpyxl's
    read-only/data-only mode.  Every input dimension is bounded so a malformed
    or hostile file cannot make ingestion/preview allocate unbounded memory.
    Limits can be supplied per instance (useful for tests) or configured via
    ``KNOWLEDGE_MAX_PARSE_*`` environment variables.
    """

    def __init__(
        self,
        *,
        max_file_bytes: int | None = None,
        max_rows: int | None = None,
        max_columns: int | None = None,
        max_cell_chars: int | None = None,
        max_text_chars: int | None = None,
        max_sheets: int | None = None,
    ) -> None:
        def resolve(value: int | None, names: tuple[str, ...], default: int) -> int:
            if value is not None:
                if value <= 0:
                    raise ValueError("parser limits must be positive")
                return value
            return _env_limit(names, default)

        self.max_file_bytes = resolve(
            max_file_bytes,
            ("KNOWLEDGE_MAX_PARSE_FILE_BYTES", "MAX_PARSE_FILE_BYTES"),
            DEFAULT_MAX_PARSE_FILE_BYTES,
        )
        self.max_rows = resolve(
            max_rows,
            ("KNOWLEDGE_MAX_PARSE_ROWS", "MAX_PARSE_ROWS"),
            DEFAULT_MAX_PARSE_ROWS,
        )
        self.max_columns = resolve(
            max_columns,
            ("KNOWLEDGE_MAX_PARSE_COLUMNS", "MAX_PARSE_COLUMNS"),
            DEFAULT_MAX_PARSE_COLUMNS,
        )
        self.max_cell_chars = resolve(
            max_cell_chars,
            ("KNOWLEDGE_MAX_PARSE_CELL_CHARS", "MAX_PARSE_CELL_CHARS"),
            DEFAULT_MAX_PARSE_CELL_CHARS,
        )
        self.max_text_chars = resolve(
            max_text_chars,
            ("KNOWLEDGE_MAX_PARSE_TEXT_CHARS", "MAX_PARSE_TEXT_CHARS"),
            DEFAULT_MAX_PARSE_TEXT_CHARS,
        )
        self.max_sheets = resolve(
            max_sheets,
            ("KNOWLEDGE_MAX_PARSE_SHEETS", "MAX_PARSE_SHEETS"),
            DEFAULT_MAX_PARSE_SHEETS,
        )

    def parse(self, path: Path) -> str:
        path = Path(path)
        suffix = path.suffix.lower()
        if suffix not in {
            ".txt",
            ".md",
            ".markdown",
            ".pdf",
            ".docx",
            ".csv",
            ".xlsx",
        }:
            raise AppError("UNSUPPORTED_DOCUMENT", f"暂不支持 {suffix or '未知'} 文件", 400)
        self._validate_path(path)
        try:
            if suffix in {".txt", ".md", ".markdown"}:
                return self._bounded_text(
                    path.read_text(encoding="utf-8", errors="ignore")
                )
            if suffix == ".pdf":
                return self._parse_pdf(path)
            if suffix == ".docx":
                return self._parse_docx(path)
            if suffix == ".csv":
                return self._parse_csv(path)
            return self._parse_xlsx(path)
        except AppError:
            raise
        except (OSError, ValueError, csv.Error) as exc:
            raise AppError(
                "DOCUMENT_PARSE_FAILED", "文档解析失败，请检查文件内容", 422
            ) from exc
        except Exception as exc:  # noqa: BLE001 - third-party parser failures
            raise AppError(
                "DOCUMENT_PARSE_FAILED", "文档解析失败，请检查文件内容", 422
            ) from exc

    def _validate_path(self, path: Path) -> None:
        if path.is_symlink() or not path.is_file():
            raise AppError("FILE_NOT_FOUND", "源文件不存在", 404)
        try:
            size = path.stat().st_size
        except OSError as exc:
            raise AppError("FILE_READ_ERROR", "文件读取失败", 500) from exc
        if size > self.max_file_bytes:
            raise AppError(
                "DOCUMENT_TOO_LARGE",
                f"文档超过 {self.max_file_bytes} 字节解析限制",
                413,
                {"maxBytes": self.max_file_bytes},
            )

    def _bounded_text(self, text: str) -> str:
        if len(text) > self.max_text_chars:
            raise AppError(
                "DOCUMENT_TEXT_TOO_LARGE",
                f"文档文本超过 {self.max_text_chars} 字符解析限制",
                413,
                {"maxChars": self.max_text_chars},
            )
        return text

    def _bounded_rows(self, rows: Iterable[Iterable[str]], *, prefix: str = "") -> str:
        lines: list[str] = []
        total_chars = len(prefix)
        if prefix:
            lines.append(prefix)
        for row in rows:
            line = "\t".join(row)
            total_chars += len(line) + 1
            if total_chars > self.max_text_chars:
                raise AppError(
                    "DOCUMENT_TEXT_TOO_LARGE",
                    f"文档文本超过 {self.max_text_chars} 字符解析限制",
                    413,
                    {"maxChars": self.max_text_chars},
                )
            lines.append(line)
        return "\n".join(lines)

    def _normalize_cell(self, value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, (datetime, date, time)):
            text = value.isoformat()
        else:
            text = str(value)
        if len(text) > self.max_cell_chars:
            raise AppError(
                "DOCUMENT_CELL_TOO_LARGE",
                f"单元格超过 {self.max_cell_chars} 字符限制",
                413,
                {"maxCellChars": self.max_cell_chars},
            )
        return text.strip()

    def _parse_pdf(self, path: Path) -> str:
        parts: list[str] = []
        total = 0
        for page in PdfReader(str(path)).pages:
            text = page.extract_text() or ""
            total += len(text) + 1
            if total > self.max_text_chars:
                raise AppError(
                    "DOCUMENT_TEXT_TOO_LARGE",
                    f"文档文本超过 {self.max_text_chars} 字符解析限制",
                    413,
                    {"maxChars": self.max_text_chars},
                )
            parts.append(text)
        return "\n".join(parts)

    def _parse_docx(self, path: Path) -> str:
        from docx import Document

        parts: list[str] = []
        total = 0
        for paragraph in Document(str(path)).paragraphs:
            text = paragraph.text
            total += len(text) + 1
            if total > self.max_text_chars:
                raise AppError(
                    "DOCUMENT_TEXT_TOO_LARGE",
                    f"文档文本超过 {self.max_text_chars} 字符解析限制",
                    413,
                    {"maxChars": self.max_text_chars},
                )
            parts.append(text)
        return "\n".join(parts)

    def _parse_csv(self, path: Path) -> str:
        raw = path.read_bytes()
        decoded: str | None = None
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                decoded = raw.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if decoded is None:
            raise AppError("DOCUMENT_ENCODING_UNSUPPORTED", "CSV 文件编码不受支持", 422)
        sample = decoded[:8192]
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
        except csv.Error:
            dialect = csv.excel
        # csv.field_size_limit is process-global.  Hold a short lock while
        # parsing so concurrent uploads cannot observe another request's limit.
        with _CSV_FIELD_LIMIT_LOCK:
            previous_limit = csv.field_size_limit(self.max_cell_chars)
            try:
                reader = csv.reader(io.StringIO(decoded, newline=""), dialect)

                def rows() -> Iterator[list[str]]:
                    for row_number, raw_row in enumerate(reader, 1):
                        if row_number > self.max_rows:
                            raise AppError(
                                "DOCUMENT_ROWS_TOO_LARGE",
                                f"CSV 行数超过 {self.max_rows} 行限制",
                                413,
                                {"maxRows": self.max_rows},
                            )
                        if len(raw_row) > self.max_columns:
                            raise AppError(
                                "DOCUMENT_COLUMNS_TOO_LARGE",
                                f"CSV 列数超过 {self.max_columns} 列限制",
                                413,
                                {"maxColumns": self.max_columns},
                            )
                        yield [self._normalize_cell(value) for value in raw_row]

                return self._bounded_rows(rows())
            except csv.Error as exc:
                raise AppError(
                    "DOCUMENT_CELL_TOO_LARGE",
                    f"单元格超过 {self.max_cell_chars} 字符限制",
                    413,
                    {"maxCellChars": self.max_cell_chars},
                ) from exc
            finally:
                csv.field_size_limit(previous_limit)

    def _parse_xlsx(self, path: Path) -> str:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        try:
            if len(workbook.worksheets) > self.max_sheets:
                raise AppError(
                    "DOCUMENT_SHEETS_TOO_LARGE",
                    f"工作表数量超过 {self.max_sheets} 张限制",
                    413,
                    {"maxSheets": self.max_sheets},
                )

            def rows() -> Iterator[list[str]]:
                row_count = 0
                for worksheet in workbook.worksheets:
                    if worksheet.max_row and worksheet.max_row > self.max_rows:
                        raise AppError(
                            "DOCUMENT_ROWS_TOO_LARGE",
                            f"工作表行数超过 {self.max_rows} 行限制",
                            413,
                            {"maxRows": self.max_rows, "sheet": worksheet.title},
                        )
                    if worksheet.max_column and worksheet.max_column > self.max_columns:
                        raise AppError(
                            "DOCUMENT_COLUMNS_TOO_LARGE",
                            f"工作表列数超过 {self.max_columns} 列限制",
                            413,
                            {"maxColumns": self.max_columns, "sheet": worksheet.title},
                        )
                    yield [f"[工作表: {self._normalize_cell(worksheet.title)}]"]
                    for row in worksheet.iter_rows(values_only=True):
                        row_count += 1
                        if row_count > self.max_rows:
                            raise AppError(
                                "DOCUMENT_ROWS_TOO_LARGE",
                                f"工作簿行数超过 {self.max_rows} 行限制",
                                413,
                                {"maxRows": self.max_rows},
                            )
                        if len(row) > self.max_columns:
                            raise AppError(
                                "DOCUMENT_COLUMNS_TOO_LARGE",
                                f"工作表列数超过 {self.max_columns} 列限制",
                                413,
                                {"maxColumns": self.max_columns, "sheet": worksheet.title},
                            )
                        yield [self._normalize_cell(value) for value in row]

            return self._bounded_rows(rows())
        finally:
            workbook.close()


class TextChunker:
    def __init__(self, chunk_size: int = 800, overlap: int = 100):
        if overlap >= chunk_size:
            raise ValueError("overlap must be smaller than chunk_size")
        self.chunk_size = chunk_size
        self.overlap = overlap

    def split(self, text: str) -> list[str]:
        normalized = re.sub(r"\r\n?", "\n", text).strip()
        if not normalized:
            return []
        chunks: list[str] = []
        start = 0
        while start < len(normalized):
            end = min(start + self.chunk_size, len(normalized))
            if end < len(normalized):
                boundary = max(normalized.rfind(mark, start, end) for mark in ("\n", "。", "！", "？"))
                if boundary > start + self.chunk_size // 2:
                    end = boundary + 1
            chunks.append(normalized[start:end].strip())
            if end >= len(normalized):
                break
            start = max(start + 1, end - self.overlap)
        return [chunk for chunk in chunks if chunk]


class KnowledgeService:
    def __init__(
        self,
        parser: DocumentParser | None = None,
        chunker: TextChunker | None = None,
        vector_indexer: VectorIndexer | None = None,
    ):
        self.parser = parser or DocumentParser()
        self.chunker = chunker or TextChunker()
        self.vector_indexer = vector_indexer

    def create_base(
        self, db: Session, *, owner_id: int, name: str, description: str | None = None
    ) -> KnowledgeBase:
        """创建知识库。

        owner_id 是「商家数据 owner」（resolve_owner 结果），与操作者
        （actor）分离：组织成员（operator）创建的知识库归属组织 owner，
        uploader 类字段才记录操作者本人。
        """
        item = KnowledgeBase(owner_id=owner_id, name=name, description=description)
        db.add(item)
        db.commit()
        db.refresh(item)
        return item

    def list_bases(self, db: Session, owner_id: int) -> list[KnowledgeBase]:
        """列出 owner_id（商家数据 owner）名下的知识库。"""
        return list(
            db.scalars(
                select(KnowledgeBase)
                .where(KnowledgeBase.owner_id == owner_id)
                .order_by(KnowledgeBase.created_at.desc())
            )
        )

    def require_owned_base(self, db: Session, base_id: int, owner_id: int) -> KnowledgeBase:
        """归属校验：base 必须属于 owner_id（商家数据 owner，resolve_owner 结果）。

        跨商家 base 一律抛 404（KNOWLEDGE_BASE_NOT_FOUND），语义与
        Chat 侧 _assert_knowledge_base_access 的 403 一致：不暴露他人知识库存在性。
        """
        item = db.get(KnowledgeBase, base_id)
        if not item or item.owner_id != owner_id:
            raise AppError("KNOWLEDGE_BASE_NOT_FOUND", "知识库不存在", 404)
        return item

    def create_document(
        self,
        db: Session,
        *,
        base_id: int,
        uploader_id: int,
        owner_id: int,
        filename: str,
        storage_path: str,
        file_size: int,
    ) -> KnowledgeDocument:
        """创建文档记录（不解析/不摄取，状态 pending）。

        owner_id 是「商家数据 owner」（resolve_owner 结果），仅用于归属校验
        （require_owned_base）；uploader_id 记录实际操作者（actor），两者
        可不同：组织成员（operator）上传文档时 owner=组织 owner、
        uploader=成员自身。
        """
        self.require_owned_base(db, base_id, owner_id)
        item = KnowledgeDocument(
            knowledge_base_id=base_id,
            uploader_id=uploader_id,
            filename=filename,
            file_type=Path(filename).suffix.lower().lstrip("."),
            storage_path=storage_path,
            file_size=file_size,
        )
        db.add(item)
        db.commit()
        db.refresh(item)
        return item

    def ingest_document(self, db: Session, document_id: int) -> KnowledgeDocument:
        document = db.get(KnowledgeDocument, document_id)
        if not document:
            raise AppError("DOCUMENT_NOT_FOUND", "文档不存在", 404)
        document.status = "processing"
        document.error_message = None
        db.commit()
        try:
            chunks = self.chunker.split(self.parser.parse(Path(document.storage_path)))
            if not chunks:
                raise ValueError("文档没有可索引的文本")
            db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_id == document.id))
            chunk_rows = [
                KnowledgeChunk(
                    knowledge_base_id=document.knowledge_base_id,
                    document_id=document.id,
                    position=position,
                    content=content,
                )
                for position, content in enumerate(chunks)
            ]
            db.add_all(chunk_rows)
            db.flush()
            vector_chunks = [
                (chunk.id, chunk.position, chunk.content) for chunk in chunk_rows
            ]
            base = db.get(KnowledgeBase, document.knowledge_base_id)
            document.source_kind = infer_source_kind(
                document.filename, base.name if base is not None else None
            )
            db.commit()

            # parse + chunk 写库成功 → keyword 检索立即可用：
            # status=indexed 且 vector_indexed=False（即使后续向量索引失败也保持可用）。
            document.status = "indexed"
            document.vector_indexed = False
            document.error_message = None
            db.commit()

            if self.vector_indexer is not None:
                try:
                    # Remove every previous generation before inserting the new DB-backed chunk IDs.
                    # This prevents shorter re-chunks from leaving searchable stale vectors.
                    asyncio.run(self.vector_indexer.store.delete_document(document.id))
                    # 向量按「文档所属 base 的商家数据 owner」归档（与上传者
                    # uploader_id 无关）：组织成员上传的文档向量归属组织 owner，
                    # 保证 Chat 按 data_owner_id 检索时能命中。
                    asyncio.run(
                        self.vector_indexer.index(
                            owner_id=(
                                base.owner_id if base is not None else document.uploader_id
                            ),
                            knowledge_base_id=document.knowledge_base_id,
                            document_id=document.id,
                            source=document.filename,
                            chunks=vector_chunks,
                        )
                    )
                    document.vector_indexed = True
                    document.error_message = None
                    db.commit()
                except Exception as exc:  # noqa: BLE001
                    # 向量索引失败 ≠ 文档不可用：降级为关键词检索，status 保持 indexed。
                    # 尽力清理本次写入的残留向量（清理失败不影响降级结果）。
                    db.rollback()
                    document = db.get(KnowledgeDocument, document_id)
                    try:
                        asyncio.run(
                            self.vector_indexer.store.delete_document(document.id)
                        )
                    except Exception:  # noqa: BLE001
                        pass
                    document.vector_indexed = False
                    document.status = "indexed"
                    document.error_message = (
                        f"向量索引失败，当前已降级使用关键词检索：{exc}"
                    )
                    db.commit()
        except Exception as exc:
            # 只有 parse 失败/文档为空/chunk 为空/DB 写入失败才置 failed
            db.rollback()
            document = db.get(KnowledgeDocument, document_id)
            document.status = "failed"
            document.error_message = str(exc)
            db.commit()
        db.refresh(document)
        return document

    def list_documents(
        self, db: Session, base_id: int, owner_id: int
    ) -> list[KnowledgeDocument]:
        """列出 base 的文档；base 必须属于 owner_id（商家数据 owner）。"""
        self.require_owned_base(db, base_id, owner_id)
        return list(
            db.scalars(
                select(KnowledgeDocument)
                .where(KnowledgeDocument.knowledge_base_id == base_id)
                .order_by(KnowledgeDocument.created_at.desc())
            )
        )
